import re
import os
from docx import Document  # DOCX 파일 처리
import openpyxl  # XLSX 파일 처리

# 🔹 개인정보 마스킹 함수
def mask_ssn(text):
    return re.sub(r'(\d{6})[-](\d{7})', r'\1-*******', text)

def mask_phone(text):
    return re.sub(r'(\d{3})-(\d{4})-(\d{4})', r'\1-****-\3', text)

def mask_email(text):
    return re.sub(r'([a-zA-Z0-9_.+-])([a-zA-Z0-9_.+-]*)@([a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+)', r'\1***@\3', text)

def mask_credit_card(text):
    return re.sub(r'(\d{4})-(\d{4})-(\d{4})-(\d{4})', r'\1-****-****-\4', text)

# 🔹 전체 마스킹 적용
def mask_sensitive_data(text):
    text = mask_ssn(text)
    text = mask_phone(text)
    text = mask_email(text)
    text = mask_credit_card(text)
    return text

# 🔹 TXT 파일 마스킹
def mask_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        masked_content = mask_sensitive_data(content)
        
        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(masked_content)
        
        return f"{file_path} (TXT) 마스킹 완료"
    except Exception as e:
        return f"TXT 마스킹 오류: {e}"

# 🔹 DOCX 파일 마스킹
def mask_docx(file_path):
    try:
        doc = Document(file_path)
        
        # 문서 내 모든 문단 마스킹 적용
        for para in doc.paragraphs:
            para.text = mask_sensitive_data(para.text)
        
        # 표 내 데이터도 마스킹
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = mask_sensitive_data(cell.text)
        
        doc.save(file_path)
        return f"{file_path} (DOCX) 마스킹 완료"
    except Exception as e:
        return f"DOCX 마스킹 오류: {e}"

# 🔹 XLSX 파일 마스킹
def mask_xlsx(file_path):
    try:
        wb = openpyxl.load_workbook(file_path)
        
        # 모든 시트 검사
        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        cell.value = mask_sensitive_data(cell.value)
        
        wb.save(file_path)
        return f"{file_path} (XLSX) 마스킹 완료"
    except Exception as e:
        return f"XLSX 마스킹 오류: {e}"

# 🔹 파일 유형별 마스킹 실행(txt, doc, xlsx)
def mask_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".txt":
        return mask_txt(file_path)
    elif ext == ".docx":
        return mask_docx(file_path)
    elif ext == ".xlsx":
        return mask_xlsx(file_path)
    else:
        return f"{file_path}: 지원되지 않는 파일 형식"

